#! /usr/bin/env python3
'''
Implementation of the "ReportFormat" module as a DOCX file
'''
from ReportFormat import ReportFormat
import ViReport_GlobalContext as GC
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from os.path import isfile
from PIL import Image
GC.report_out_doc = None

def list_number(doc, par, prev=None, level=None, num=True):
    """
    Makes a paragraph into a list item with a specific level and
    optional restart.

    An attempt will be made to retreive an abstract numbering style that
    corresponds to the style of the paragraph. If that is not possible,
    the default numbering or bullet style will be used based on the
    ``num`` parameter.

    Obtained from: https://stackoverflow.com/a/51830413/2134991

    Parameters
    ----------
    doc : docx.document.Document
        The document to add the list into.
    par : docx.paragraph.Paragraph
        The paragraph to turn into a list item.
    prev : docx.paragraph.Paragraph or None
        The previous paragraph in the list. If specified, the numbering
        and styles will be taken as a continuation of this paragraph.
        If omitted, a new numbering scheme will be started.
    level : int or None
        The level of the paragraph within the outline. If ``prev`` is
        set, defaults to the same level as in ``prev``. Otherwise,
        defaults to zero.
    num : bool
        If ``prev`` is :py:obj:`None` and the style of the paragraph
        does not correspond to an existing numbering style, this will
        determine wether or not the list will be numbered or bulleted.
        The result is not guaranteed, but is fairly safe for most Word
        templates.
    """
    xpath_options = {
        True: {'single': 'count(w:lvl)=1 and ', 'level': 0},
        False: {'single': '', 'level': level},
    }

    def style_xpath(prefer_single=True):
        """
        The style comes from the outer-scope variable ``par.style.name``.
        """
        style = par.style.style_id
        return (
            'w:abstractNum['
                '{single}w:lvl[@w:ilvl="{level}"]/w:pStyle[@w:val="{style}"]'
            ']/@w:abstractNumId'
        ).format(style=style, **xpath_options[prefer_single])

    def type_xpath(prefer_single=True):
        """
        The type is from the outer-scope variable ``num``.
        """
        type = 'decimal' if num else 'bullet'
        return (
            'w:abstractNum['
                '{single}w:lvl[@w:ilvl="{level}"]/w:numFmt[@w:val="{type}"]'
            ']/@w:abstractNumId'
        ).format(type=type, **xpath_options[prefer_single])

    def get_abstract_id():
        """
        Select as follows:

            1. Match single-level by style (get min ID)
            2. Match exact style and level (get min ID)
            3. Match single-level decimal/bullet types (get min ID)
            4. Match decimal/bullet in requested level (get min ID)
            3. 0
        """
        for fn in (style_xpath, type_xpath):
            for prefer_single in (True, False):
                xpath = fn(prefer_single)
                ids = numbering.xpath(xpath)
                if ids:
                    return min(int(x) for x in ids)
        return 0

    if (prev is None or
            prev._p.pPr is None or
            prev._p.pPr.numPr is None or
            prev._p.pPr.numPr.numId is None):
        if level is None:
            level = 0
        numbering = doc.part.numbering_part.numbering_definitions._numbering
        # Compute the abstract ID first by style, then by num
        anum = get_abstract_id()
        # Set the concrete numbering based on the abstract numbering ID
        num = numbering.add_num(anum)
        # Make sure to override the abstract continuation property
        num.add_lvlOverride(ilvl=level).add_startOverride(1)
        # Extract the newly-allocated concrete numbering ID
        num = num.numId
    else:
        if level is None:
            level = prev._p.pPr.numPr.ilvl.val
        # Get the previous concrete numbering ID
        num = prev._p.pPr.numPr.numId.val
    par._p.get_or_add_pPr().get_or_add_numPr().get_or_add_numId().val = num
    par._p.get_or_add_pPr().get_or_add_numPr().get_or_add_ilvl().val = level

def docx_init():
    if GC.report_out_doc is None:
        GC.report_out_doc_filename = '%s/Report.docx' % GC.OUT_DIR
        GC.report_out_doc = Document()
        GC.report_out_doc.add_heading("ViReport v%s — %s" % (GC.VIREPORT_VERSION, datetime.today().strftime('%Y-%m-%d')), 0)

class ReportFormat_DOCX(ReportFormat):
    def init():
        pass

    def finalize():
        pass

    def cite():
        return GC.CITATION_VIREPORT

    def section(s):
        docx_init(); GC.report_out_doc.add_heading(s, 1); GC.report_out_doc.add_paragraph()

    def subsection(s):
        docx_init(); GC.report_out_doc.add_heading(s, 2); GC.report_out_doc.add_paragraph()

    def write(s):
        docx_init()
        if len(GC.report_out_doc.paragraphs) == 0:
            GC.report_out_doc.add_paragraph()
        GC.report_out_doc.paragraphs[-1].add_run(s)

    def writeln(s):
        docx_init(); ReportFormat_DOCX.write(s); GC.report_out_doc.add_paragraph()

    def bullets(items, level=0):
        docx_init()
        if len(GC.report_out_doc.paragraphs[-1].text) == 0:
            del GC.report_out_doc.paragraphs[-1]
        for item in items:
            if isinstance(item, str):
                tmp = GC.report_out_doc.add_paragraph(item, style=('List Bullet %d'%(level+1)).rstrip(' 1'))
                list_number(GC.report_out_doc, tmp, level=level, num=False)
            elif isinstance(item, list):
                ReportFormat_DOCX.bullets(item, level=level+1)
            else:
                raise ValueError("Invalid bullet item type: %s" % type(item))

    def figure(filename, caption=None, width=None, height=None, keep_aspect_ratio=True):
        if not filename.startswith(GC.OUT_DIR_REPORTFILES):
            raise ValueError("Figures must be in report files directory: %s" % GC.OUT_DIR_REPORTFILES)
        if filename.lower().endswith('.pdf'):
            png_filename = '%s.%s' % ('.'.join(filename.split('.')[:-1]), 'png')
            if not isfile(png_filename):
                GC.pdf_to_png(filename, png_filename)
            filename = png_filename
        w,h = Image.open(filename).size; ar = w/h
        if width is not None:
            width = Inches(6*width)
        if height is not None:
            height = Inches(8.5*height)
        pic = GC.report_out_doc.add_picture(filename, width=width, height=height)
        if keep_aspect_ratio:
            if pic.width/pic.height > ar:
                pic.width = int(pic.height * ar)
            else:
                pic.height = int(pic.width / ar)
        GC.report_out_doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption is not None:
            GC.report_out_doc.add_paragraph(caption, style='Caption')
        GC.report_out_doc.add_paragraph()

    def close():
        docx_init(); GC.report_out_doc.save(GC.report_out_doc_filename)
        return GC.report_out_doc_filename
